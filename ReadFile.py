from docx import Document
#doc = Document("test.docx")
import logging
logger = logging.getLogger('Simple_Agent_Tool.ReadFile')

class ReadFile:
    text:str
    def __init__(self):
        logger.info('初始化读取文件工具')
        pass
    def read_by_docx(self,file_path:str):
        logger.info(f'读取docx文件，文件路径：{file_path}')
        self.text=str()
        try:
            doc=Document(file_path)
            for p in doc.paragraphs:
                if p.text=='':
                    continue
                self.text+=p.text+'\n'
            logger.info(f'读取成功')
        except Exception as e:
            logger.error(f'读取docx文件失败，文件路径：{file_path}，错误信息：{e}')
            return ''
        return self.text
    def read_by_markdown(self,file_path:str):
        logger.info(f'读取markdown文件，文件路径：{file_path}')
        self.text=str()
        try:
            with open(file_path,'r',encoding='utf-8') as f:
                self.text=f.read()
                logger.info(f'读取成功')
        except Exception as e:
            logger.error(f'读取markdown文件失败，文件路径：{file_path}，错误信息：{e}')
        return self.text
#rf=ReadFile()
#rf.read_by_docx('喜羊羊_简历.docx')
#print(rf.text)